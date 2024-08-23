import streamlit as st
from docx import Document
from docx.text.paragraph import Paragraph
import io
import zipfile
from PIL import Image
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

def read_docx(file):
    return Document(io.BytesIO(file.read()))

def correct_text_in_docx(doc, corrections):
    # コメントを保存 (新しい方法)
    comments = doc.part.comments._comments if hasattr(doc.part, 'comments') else {}

    for paragraph in doc.paragraphs:
        correct_paragraph(paragraph, corrections)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    correct_paragraph(paragraph, corrections)

    # コメントを復元 (新しい方法)
    if hasattr(doc.part, 'comments'):
        doc.part.comments._comments = comments

    return doc

def correct_paragraph(paragraph, corrections):
    if not paragraph.runs:
        return

    # 段落全体のテキストを取得し、修正を適用
    full_text = paragraph.text
    original_text = full_text
    for wrong, correct in corrections.items():
        full_text = full_text.replace(wrong, correct)

    # テキストが変更された場合のみ処理を行う
    if full_text != original_text:
        # 元の段落のスタイルとコメントを保存
        original_style = paragraph.style
        original_comments = get_paragraph_comments(paragraph)

        # 段落をクリアする
        paragraph.clear()

        # 新しいテキストを追加し、元の書式を可能な限り保持する
        start = 0
        for run in paragraph.runs:
            end = min(start + len(run.text), len(full_text))
            if start < end:
                new_run = paragraph.add_run(full_text[start:end])
                copy_run_format(run, new_run)
            start = end

        # 残りのテキストがあれば追加
        if start < len(full_text):
            paragraph.add_run(full_text[start:])

        # 段落のスタイルを復元
        paragraph.style = original_style

        # コメントを復元
        restore_paragraph_comments(paragraph, original_comments)

def get_paragraph_comments(paragraph):
    comments = []
    for run in paragraph.runs:
        comment_reference = run._r.find(qn('w:commentReference'))
        if comment_reference is not None:
            comment_id = comment_reference.get(qn('w:id'))
            if hasattr(paragraph.part, 'comments'):
                comment = paragraph.part.comments.get_comment_by_id(comment_id)
                if comment:
                    comments.append((run, comment))
    return comments

def restore_paragraph_comments(paragraph, comments):
    if hasattr(paragraph.part, 'comments'):
        for run, comment in comments:
            new_run = paragraph.add_run()
            new_run._r.append(comment.reference)
            paragraph.part.comments._comments[comment._id] = comment

def copy_run_format(src_run, dest_run):
    # 既存の属性をコピー
    dest_run.bold = src_run.bold
    dest_run.italic = src_run.italic
    dest_run.underline = src_run.underline
    dest_run.font.name = src_run.font.name
    dest_run.font.size = src_run.font.size
    dest_run.font.color.rgb = src_run.font.color.rgb
    
    # 追加の属性をコピー
    dest_run.style = src_run.style
    dest_run.font.highlight_color = src_run.font.highlight_color
    dest_run.font.subscript = src_run.font.subscript
    dest_run.font.superscript = src_run.font.superscript
    dest_run.font.strike = src_run.font.strike
    dest_run.font.shadow = src_run.font.shadow
    dest_run.font.outline = src_run.font.outline
    dest_run.font.rtl = src_run.font.rtl
    dest_run.font.small_caps = src_run.font.small_caps
    dest_run.font.cs_bold = src_run.font.cs_bold
    dest_run.font.cs_italic = src_run.font.cs_italic
    
    # 既存のコードに加えて以下を追加
    dest_run.style = src_run.style
    dest_run._element.rPr.append(src_run._element.rPr)

def correct_text_for_display(text, corrections):
    for wrong, correct in corrections.items():
        text = text.replace(wrong, f'<span style="color:red;">{correct}</span>')
    return text

def correct_text_for_download(text, corrections):
    for wrong, correct in corrections.items():
        text = text.replace(wrong, correct)
    return text

st.title('AI文例作成サポート用表記揺れチェック') 
st.markdown("""
表記揺れチェック＆修正したファイル（docxかmd）をアップロードしてください。  
検品したい表記を選択して、開始ボタンを押してください。  
ダウンロードファイル用のボタンが出現します（.docxファイルが出てきます）。
""")

image = Image.open('nyan_ai_writing_01.jpg')
st.image(image,use_column_width=True)

uploaded_files = st.file_uploader("WordまたはMarkdownファイルをアップロードしてください（複数可）", type=["docx", "md"], accept_multiple_files=True)

# ユーザーが選択できる表記揺れのリスト
options = {
    "下さ": "くださ",
    "頂": "いただ",
    "虫歯": "むし歯",
    "出来": "でき",
    "致し": "いたし",
    "当クリニック": "当院",
}

# ユーザーが選択した表記揺れ
selected_options = st.multiselect("修正したい表記揺れを選択してください", list(options.keys()))

# ユーザーが自由に入力できるキーワードのリスト
user_keywords = st.text_area("追加でチェックしたいキーワードを「キーワード:変換後の文字」  \n（例：内科をナイカにしたい時「内科:ナイカ」）の形式で1行ずつ入力してください。コロンは半角です").split('\n')

# 開始ボタン
if st.button('開始'):
    # ユーザーが選択した表記揺れと自由に入力したキーワードを修正するための辞書
    corrections = {key: options[key] for key in selected_options}
    corrections.update({key.split(':')[0]: key.split(':')[1] for key in user_keywords if ':' in key})

    if len(uploaded_files) == 1:
        # Only one file, so download it directly
        uploaded_file = uploaded_files[0]
        if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # This is a Word file (.docx)
            doc = read_docx(uploaded_file)
            text = '\n'.join([para.text for para in doc.paragraphs])
            
            st.write('アップロードされたファイルの内容:')
            st.write(text)

            corrected_text_for_display = correct_text_for_display(text, corrections)
            st.write('修正後のテキスト:')
            st.markdown(corrected_text_for_display, unsafe_allow_html=True)

            corrected_doc = correct_text_in_docx(doc, corrections)
            
            docx_buffer = io.BytesIO()
            corrected_doc.save(docx_buffer)
            docx_buffer.seek(0)

            st.download_button(
                label="修正後のファイルをダウンロード",
                data=docx_buffer,
                file_name='corrected.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
        else:
            st.error("Unsupported file type: " + uploaded_file.type)
    else:
        # Multiple files, so zip them
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
            for i, uploaded_file in enumerate(uploaded_files):
                if uploaded_file is not None and uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc = read_docx(uploaded_file)
                    text = '\n'.join([para.text for para in doc.paragraphs])
                    
                    st.write('アップロードされたファイルの内容:')
                    st.write(text)

                    corrected_text_for_display = correct_text_for_display(text, corrections)
                    st.write('修正後のテキスト:')
                    st.markdown(corrected_text_for_display, unsafe_allow_html=True)

                    corrected_doc = correct_text_in_docx(doc, corrections)
                    
                    docx_buffer = io.BytesIO()
                    corrected_doc.save(docx_buffer)
                    docx_buffer.seek(0)
                    
                    zip_file.writestr(f'corrected_{i}.docx', docx_buffer.getvalue())
                else:
                    st.error("Unsupported file type: " + uploaded_file.type)

        # Finish the zip file
        zip_file.close()
        zip_buffer.seek(0)

        # Create a download button for the zip file
        st.download_button(
            label="修正後のファイルをダウンロード",
            data=zip_buffer.getvalue(),
            file_name='corrected_files.zip',
            mime='application/zip',
        )